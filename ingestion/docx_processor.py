"""
DOCX document processor for medical texts.
"""

import os
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import re

from docx import Document
from docx.shared import Inches
from docx.document import Document as DocumentType
from docx.text.paragraph import Paragraph
from docx.table import Table

from .chunker import DocumentChunk

logger = logging.getLogger(__name__)


class DOCXProcessor:
    """Processor for DOCX medical documents."""
    
    def __init__(self):
        """Initialize DOCX processor."""
        pass
    
    def process_docx_file(self, file_path: str) -> Dict[str, Any]:
        """
        Process a DOCX file and extract text content.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted document data
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
        logger.info(f"Processing DOCX file: {file_path}")
        
        try:
            # Load DOCX document
            doc = Document(file_path)
            
            # Extract document properties
            title = self._extract_title(doc, file_path)
            content = self._extract_content(doc)
            metadata = self._extract_metadata(doc, file_path)
            
            # Count pages (approximate)
            page_count = self._estimate_page_count(content)
            
            result = {
                "title": title,
                "content": content,
                "source": file_path,
                "metadata": {
                    **metadata,
                    "file_type": "docx",
                    "estimated_pages": page_count,
                    "character_count": len(content),
                    "word_count": len(content.split())
                }
            }
            
            logger.info(f"✓ Processed DOCX: {len(content)} chars, ~{page_count} pages")
            return result
            
        except Exception as e:
            logger.error(f"Failed to process DOCX file {file_path}: {str(e)}")
            raise
    
    def _extract_title(self, doc: DocumentType, file_path: str) -> str:
        """Extract document title from DOCX."""
        # Try document properties first
        if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
            return doc.core_properties.title.strip()
        
        # Try first heading or paragraph
        for paragraph in doc.paragraphs[:5]:  # Check first 5 paragraphs
            if paragraph.text.strip():
                # If it's a heading style or short text, likely a title
                if (hasattr(paragraph.style, 'name') and 
                    ('heading' in paragraph.style.name.lower() or 
                     'title' in paragraph.style.name.lower())):
                    return paragraph.text.strip()
                
                # If short text (likely title)
                if len(paragraph.text.strip()) < 100:
                    return paragraph.text.strip()
        
        # Fallback to filename
        return Path(file_path).stem.replace('_', ' ').replace('-', ' ').title()
    
    def _extract_content(self, doc: DocumentType) -> str:
        """Extract all text content from DOCX document."""
        content_parts = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Add heading formatting for styled paragraphs
                if hasattr(paragraph.style, 'name'):
                    style_name = paragraph.style.name.lower()
                    if 'heading' in style_name:
                        if 'heading 1' in style_name:
                            text = f"\n# {text}\n"
                        elif 'heading 2' in style_name:
                            text = f"\n## {text}\n"
                        elif 'heading 3' in style_name:
                            text = f"\n### {text}\n"
                        else:
                            text = f"\n**{text}**\n"
                
                content_parts.append(text)
        
        # Extract tables
        for table in doc.tables:
            table_text = self._extract_table_content(table)
            if table_text:
                content_parts.append(f"\n[TABELLA]\n{table_text}\n[/TABELLA]\n")
        
        # Join all content
        full_content = '\n'.join(content_parts)
        
        # Clean up excessive whitespace
        full_content = re.sub(r'\n\s*\n\s*\n', '\n\n', full_content)
        full_content = re.sub(r' +', ' ', full_content)
        
        return full_content.strip()
    
    def _extract_table_content(self, table: Table) -> str:
        """Extract text from a table."""
        table_text = []
        
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            
            if row_text:
                table_text.append(' | '.join(row_text))
        
        return '\n'.join(table_text)
    
    def _extract_metadata(self, doc: DocumentType, file_path: str) -> Dict[str, Any]:
        """Extract metadata from DOCX document."""
        metadata = {}
        
        # Core properties
        core_props = doc.core_properties
        if hasattr(core_props, 'author') and core_props.author:
            metadata['author'] = core_props.author
        if hasattr(core_props, 'subject') and core_props.subject:
            metadata['subject'] = core_props.subject
        if hasattr(core_props, 'created') and core_props.created:
            metadata['created'] = core_props.created.isoformat()
        if hasattr(core_props, 'modified') and core_props.modified:
            metadata['modified'] = core_props.modified.isoformat()
        if hasattr(core_props, 'keywords') and core_props.keywords:
            metadata['keywords'] = core_props.keywords
        
        # Document statistics
        metadata['paragraphs_count'] = len(doc.paragraphs)
        metadata['tables_count'] = len(doc.tables)
        
        # File info
        file_stat = os.stat(file_path)
        metadata['file_size'] = file_stat.st_size
        metadata['file_path'] = file_path
        
        return metadata
    
    def _estimate_page_count(self, content: str) -> int:
        """Estimate page count based on content length."""
        # Rough estimate: ~500 words per page for medical texts
        word_count = len(content.split())
        estimated_pages = max(1, round(word_count / 500))
        return estimated_pages
    
    def process_multiple_docx_files(self, folder_path: str) -> List[Dict[str, Any]]:
        """
        Process multiple DOCX files in a folder.
        
        Args:
            folder_path: Path to folder containing DOCX files
            
        Returns:
            List of processed document data
        """
        if not os.path.exists(folder_path):
            raise FileNotFoundError(f"Folder not found: {folder_path}")
        
        docx_files = list(Path(folder_path).glob("*.docx"))
        
        # Filter out temporary files (starting with ~$)
        docx_files = [f for f in docx_files if not f.name.startswith('~$')]
        
        if not docx_files:
            logger.warning(f"No DOCX files found in {folder_path}")
            return []
        
        logger.info(f"Found {len(docx_files)} DOCX files to process")
        
        processed_docs = []
        
        for file_path in docx_files:
            try:
                doc_data = self.process_docx_file(str(file_path))
                processed_docs.append(doc_data)
                
            except Exception as e:
                logger.error(f"Failed to process {file_path}: {str(e)}")
                continue
        
        logger.info(f"Successfully processed {len(processed_docs)}/{len(docx_files)} DOCX files")
        return processed_docs


# Factory function
def create_docx_processor() -> DOCXProcessor:
    """Create DOCX processor instance."""
    return DOCXProcessor()


# Example usage
if __name__ == "__main__":
    processor = create_docx_processor()
    
    # Example processing single file
    # doc_data = processor.process_docx_file("sample_medical.docx")
    # print(f"Title: {doc_data['title']}")
    # print(f"Content length: {len(doc_data['content'])} chars")
    # print(f"Estimated pages: {doc_data['metadata']['estimated_pages']}")
    
    # Example processing folder
    # docs = processor.process_multiple_docx_files("medical_documents/")
    # print(f"Processed {len(docs)} documents")