"""
Test pipeline: ingestione di un file DOCX verso Postgres con embeddings offline.
"""

import os
import pytest
import asyncio
import tempfile

from uuid import UUID

import docx

from ingestion.ingest import DocumentIngestionPipeline
from agent.models import IngestionConfig
from agent.db_utils import db_pool

# Impostazioni ambiente per test
os.environ.setdefault("EMBEDDINGS_OFFLINE", "1")
os.environ.setdefault("DATABASE_URL", "postgresql://rag_user:rag_password@localhost:55432/rag_db")

if os.name == "nt":
    try:
        asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())
    except Exception:
        pass


@pytest.mark.asyncio
async def test_docx_ingestion_inserts_documents_and_chunks():
    # Crea un DOCX temporaneo
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_docx = os.path.join(tmpdir, "test_ingest.docx")
        d = docx.Document()
        d.add_paragraph("Titolo documento di prova")
        d.add_paragraph("Contenuto clinico: dolore spalla e tendinopatia.")
        d.save(tmp_docx)

        # Pipeline con config default e grafo disabilitato per velocità
        cfg = IngestionConfig()
        setattr(cfg, "skip_graph_building", True)

        pipeline = DocumentIngestionPipeline(config=cfg, documents_folder=tmpdir, clean_before_ingest=True)
        await pipeline.initialize()
        results = await pipeline.ingest_documents(tenant_slug="default")
        await pipeline.close()

        # Verifiche: 1 documento, chunks inseriti
        assert len(results) == 1
        assert results[0].success is True
        document_title = results[0].title

        # Query Postgres per verifiche
        async with db_pool.acquire() as conn:
            doc = await conn.fetchrow("SELECT id, title FROM documents ORDER BY created_at DESC LIMIT 1")
            assert doc is not None
            assert document_title in doc["title"]

            chunks = await conn.fetch("SELECT document_id, content, embedding FROM chunks WHERE document_id = $1", doc["id"])
            assert len(chunks) > 0
            # In offline, l'embedding è vettore di zeri di dimensione 1536 (inserita come ::vector)
            assert chunks[0]["embedding"] is not None
