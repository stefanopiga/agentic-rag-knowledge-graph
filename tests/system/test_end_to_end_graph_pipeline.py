"""
Test E2E pipeline completa con grafo abilitato e embeddings offline.
"""

import os
import pytest
import asyncio
import tempfile
from uuid import UUID

from ingestion.ingest import DocumentIngestionPipeline
from agent.models import IngestionConfig

# Skip se mancano le variabili per DB/Neo4j/LLM/Embedding
if not (
    os.getenv("DATABASE_URL") and os.getenv("NEO4J_URI") and os.getenv("NEO4J_PASSWORD") and os.getenv("LLM_API_KEY") and os.getenv("EMBEDDING_API_KEY")
):
    pytest.skip("Ambiente cloud non configurato: DATABASE_URL/NEO4J_URI/NEO4J_PASSWORD/LLM_API_KEY/EMBEDDING_API_KEY mancanti", allow_module_level=True)

# Imposta offline embeddings
os.environ.setdefault("EMBEDDINGS_OFFLINE", "1")

if os.name == "nt":
    try:
        asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())
    except Exception:
        pass


@pytest.mark.asyncio
async def test_end_to_end_ingestion_with_graph():
    with tempfile.TemporaryDirectory() as tmpdir:
        # Crea documento di prova
        sample_path = os.path.join(tmpdir, "case_clinico.docx")
        try:
            import docx
            d = docx.Document()
            d.add_paragraph("Caso clinico: spalla")
            d.add_paragraph("Tendinopatia sovraspinato")
            d.save(sample_path)
        except Exception:
            pytest.skip("python-docx non disponibile")

        cfg = IngestionConfig()
        setattr(cfg, "skip_graph_building", False)

        pipeline = DocumentIngestionPipeline(config=cfg, documents_folder=tmpdir, clean_before_ingest=True)
        await pipeline.initialize()
        results = await pipeline.ingest_documents(tenant_slug="default")
        await pipeline.close()

        assert len(results) == 1
        assert results[0].success is True
        assert results[0].relationships_created >= 1  # almeno 1 episodio nel grafo
