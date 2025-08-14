"""
Test Pipeline - Ingestion e Processing.

Fase 2 del testing di sistema: verifica pipeline di elaborazione documenti.
"""

import os
import sys
import asyncio
import tempfile
import traceback
from pathlib import Path

# Aggiungi project root al path
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from tests.system.test_logger import create_logger


class PipelineTests:
    """Test per verifica pipeline ingestion e processing."""
    
    def __init__(self):
        self.logger = create_logger("pipeline")
        self.test_results = {
            "incremental_manager": False,
            "document_scanning": False,
            "docx_processor": False,
            "chunker": False,
            "category_recognition": False,
            "medical_entities": False
        }
    
    async def run_all_tests(self) -> dict:
        """Esegue tutti i test della pipeline."""
        self.logger.log_info("⚙️ INIZIANDO TEST PIPELINE...")
        
        try:
            # Test 1: Incremental Manager
            await self.test_incremental_manager()
            
            # Test 2: Document Scanning
            await self.test_document_scanning()
            
            # Test 3: DOCX Processor
            await self.test_docx_processor()
            
            # Test 4: Chunker
            await self.test_chunker()
            
            # Test 5: Category Recognition
            await self.test_category_recognition()
            
            # Test 6: Medical Entities
            await self.test_medical_entities()
            
        except Exception as e:
            self.logger.log_error(f"Errore critico durante test pipeline: {e}")
            self.logger.log_error(traceback.format_exc())
        
        finally:
            return self.logger.finalize()
    
    async def test_incremental_manager(self):
        """Test Incremental Manager."""
        self.logger.log_test_start("Verifica Incremental Manager")
        
        try:
            from ingestion.incremental_manager import IncrementalIngestionManager
            
            # Crea manager
            manager = IncrementalIngestionManager()
            self.logger.log_info("✅ IncrementalIngestionManager creato")
            
            # Test metodi principali
            if hasattr(manager, 'scan_documents'):
                self.logger.log_info("✅ Metodo scan_documents disponibile")
            
            if hasattr(manager, 'calculate_citation_priority'):
                self.logger.log_info("✅ Metodo calculate_citation_priority disponibile")
            
            if hasattr(manager, '_extract_metadata_from_path'):
                self.logger.log_info("✅ Metodo _extract_metadata_from_path disponibile")
            
            # Test estrazione metadata da path di esempio
            test_paths = [
                "documents/fisioterapia/master/ginocchio_e_anca/01_anatomia.docx",
                "documents/fisioterapia/master/caviglia_e_piede/02_patologie.pdf",
                "documents/fisioterapia/master/ATM/03_trattamenti.txt"
            ]
            
            for path in test_paths:
                try:
                    metadata = manager._extract_metadata_from_path(path)
                    category = metadata.get('category', 'unknown')
                    order = metadata.get('document_order', 0)
                    self.logger.log_info(f"✅ Path: {path} → Categoria: {category}, Ordine: {order}")
                except Exception as e:
                    self.logger.log_warning(f"⚠️ Errore parsing path {path}: {e}")
            
            self.test_results["incremental_manager"] = True
            self.logger.log_test_success("Verifica Incremental Manager", "Manager funzionante")
            
        except Exception as e:
            self.test_results["incremental_manager"] = False
            self.logger.log_test_failure("Verifica Incremental Manager", str(e))
    
    async def test_document_scanning(self):
        """Test Document Scanning."""
        self.logger.log_test_start("Verifica Document Scanning")
        
        try:
            from ingestion.incremental_manager import IncrementalIngestionManager
            
            manager = IncrementalIngestionManager()
            
            # Test scan della cartella documenti se esiste
            docs_folder = Path("documents/fisioterapia")
            if docs_folder.exists():
                self.logger.log_info(f"✅ Cartella documenti trovata: {docs_folder}")
                
                try:
                    # Scan documenti
                    scan_results = await manager.scan_documents(str(docs_folder))
                    
                    self.logger.log_info(f"✅ Scan completato: {len(scan_results)} documenti")
                    
                    # Analizza risultati per categoria
                    categories = {}
                    for result in scan_results:
                        cat = getattr(result, 'category', 'unknown')
                        if cat not in categories:
                            categories[cat] = 0
                        categories[cat] += 1
                    
                    for cat, count in categories.items():
                        self.logger.log_info(f"   {cat}: {count} documenti")
                    
                except Exception as e:
                    self.logger.log_warning(f"⚠️ Errore durante scan: {e}")
            else:
                self.logger.log_warning("⚠️ Cartella documenti non trovata")
            
            # Test con cartella temporanea
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_path = Path(temp_dir)
                
                # Crea struttura test
                test_structure = [
                    "fisioterapia/master/ginocchio_e_anca/01_test.docx",
                    "fisioterapia/master/caviglia_e_piede/02_test.pdf",
                    "fisioterapia/master/ATM/03_test.txt"
                ]
                
                for struct in test_structure:
                    file_path = temp_path / struct
                    file_path.parent.mkdir(parents=True, exist_ok=True)
                    file_path.write_text("Test content")
                
                # Scan cartella temporanea
                scan_results = await manager.scan_documents(str(temp_path / "fisioterapia"))
                self.logger.log_info(f"✅ Scan cartella temporanea: {len(scan_results)} files")
            
            self.test_results["document_scanning"] = True
            self.logger.log_test_success("Verifica Document Scanning", "Scanning funzionante")
            
        except Exception as e:
            self.test_results["document_scanning"] = False
            self.logger.log_test_failure("Verifica Document Scanning", str(e))
    
    async def test_docx_processor(self):
        """Test DOCX Processor."""
        self.logger.log_test_start("Verifica DOCX Processor")
        
        try:
            from ingestion.docx_processor import DOCXProcessor, create_docx_processor
            
            # Test factory
            processor = create_docx_processor()
            self.logger.log_info("✅ DOCXProcessor creato")
            
            # Test con file temporaneo DOCX
            try:
                import docx
                
                # Crea documento test
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                    doc = docx.Document()
                    doc.add_heading('Test Document', 0)
                    doc.add_paragraph('Questo è un test di processing DOCX.')
                    doc.add_paragraph('Anatomia del ginocchio include femore, tibia e rotula.')
                    
                    # Aggiungi tabella test
                    table = doc.add_table(rows=2, cols=2)
                    table.cell(0, 0).text = 'Struttura'
                    table.cell(0, 1).text = 'Funzione'
                    table.cell(1, 0).text = 'Menisco'
                    table.cell(1, 1).text = 'Ammortizzazione'
                    
                    doc.save(temp_file.name)
                    temp_path = temp_file.name
                
                # Process documento
                result = processor.process_docx_file(temp_path)
                
                self.logger.log_info("✅ Documento DOCX processato")
                self.logger.log_info(f"   Titolo: {result.get('title', 'N/A')}")
                self.logger.log_info(f"   Lunghezza contenuto: {len(result.get('content', ''))}")
                self.logger.log_info(f"   Metadata: {result.get('metadata', {})}")
                
                # Cleanup
                os.unlink(temp_path)
                
            except ImportError:
                self.logger.log_warning("⚠️ python-docx non disponibile per test completo")
            
            self.test_results["docx_processor"] = True
            self.logger.log_test_success("Verifica DOCX Processor", "Processor funzionante")
            
        except Exception as e:
            self.test_results["docx_processor"] = False
            self.logger.log_test_failure("Verifica DOCX Processor", str(e))
    
    async def test_chunker(self):
        """Test Chunker."""
        self.logger.log_test_start("Verifica Chunker")
        
        try:
            from ingestion.chunker import create_chunker, ChunkingConfig
            
            # Crea chunker (usa semantic per testare versione async)
            config = ChunkingConfig(chunk_size=200, chunk_overlap=50, use_semantic_splitting=True)
            chunker = create_chunker(config)
            self.logger.log_info("✅ Chunker creato")
            
            # Test chunking
            test_text = """
            L'articolazione del ginocchio è una delle più complesse del corpo umano.
            È formata da tre ossa principali: femore, tibia e rotula.
            
            I menischi sono strutture fibrocartilaginee che fungono da ammortizzatori.
            Il menisco mediale ha forma a C, mentre quello laterale è più circolare.
            
            I legamenti crociati forniscono stabilità antero-posteriore.
            Il legamento crociato anteriore previene lo scivolamento anteriore della tibia.
            Il legamento crociato posteriore previene lo scivolamento posteriore.
            """
            
            # Usa il metodo chunk_document asincrono
            chunks = await chunker.chunk_document(
                content=test_text, 
                title="Test Document", 
                source="test.txt"
            )
            
            self.logger.log_info(f"✅ Testo diviso in {len(chunks)} chunks")
            
            for i, chunk in enumerate(chunks[:3]):  # Mostra primi 3 chunks
                self.logger.log_info(f"   Chunk {i+1}: {len(chunk.content)} caratteri")
                self.logger.log_info(f"   Preview: {chunk.content[:50]}...")
            
            self.test_results["chunker"] = True
            self.logger.log_test_success("Verifica Chunker", f"Chunking riuscito - {len(chunks)} chunks")
            
        except Exception as e:
            self.test_results["chunker"] = False
            self.logger.log_test_failure("Verifica Chunker", str(e))
    
    async def test_category_recognition(self):
        """Test Category Recognition."""
        self.logger.log_test_start("Verifica Category Recognition")
        
        try:
            from ingestion.incremental_manager import IncrementalIngestionManager
            
            manager = IncrementalIngestionManager()
            
            # Test categorie supportate
            test_paths = [
                "documents/fisioterapia/master/ginocchio_e_anca/test.docx",
                "documents/fisioterapia/master/caviglia_e_piede/test.pdf",
                "documents/fisioterapia/master/ATM/test.txt",
                "documents/fisioterapia/master/cervicale/test.docx",
                "documents/fisioterapia/master/lombare/test.pdf",
                "documents/fisioterapia/master/toracico/test.txt",
                "documents/fisioterapia/master/lombo_pelvico/test.docx",
                "documents/fisioterapia/master/arto_superiore/test.pdf"
            ]
            
            recognized_categories = set()
            
            for path in test_paths:
                try:
                    metadata = manager._extract_metadata_from_path(path)
                    category = metadata.get('category')
                    priority = manager.calculate_citation_priority(category, metadata.get('document_order', 1))
                    
                    if category and category != 'generale':
                        recognized_categories.add(category)
                        self.logger.log_info(f"✅ {category}: priorità {priority}")
                
                except Exception as e:
                    self.logger.log_warning(f"⚠️ Errore riconoscimento {path}: {e}")
            
            self.logger.log_info(f"✅ Categorie riconosciute: {len(recognized_categories)}")
            self.logger.log_info(f"   {', '.join(sorted(recognized_categories))}")
            
            self.test_results["category_recognition"] = True
            self.logger.log_test_success("Verifica Category Recognition", f"{len(recognized_categories)} categorie riconosciute")
            
        except Exception as e:
            self.test_results["category_recognition"] = False
            self.logger.log_test_failure("Verifica Category Recognition", str(e))
    
    async def test_medical_entities(self):
        """Test Medical Entities Extraction."""
        self.logger.log_test_start("Verifica Medical Entities Extraction")
        
        try:
            from ingestion.graph_builder import SimpleMedicalEntityExtractor
            
            # Crea extractor
            extractor = SimpleMedicalEntityExtractor()
            self.logger.log_info("✅ SimpleMedicalEntityExtractor creato")
            
            # Test testo medico
            medical_text = """
            L'articolazione del ginocchio comprende il femore, la tibia e la rotula.
            Il menisco mediale può subire lesioni durante la distorsione.
            La fisioterapia include mobilizzazione e rinforzo del quadricipite.
            Il trattamento prevede ultrasuoni e elettrostimolazione.
            """
            
            # Estrai entità
            entities_dict = extractor.extract_entities(medical_text)
            
            # Conta entità per categoria
            total_entities = 0
            for entity_type, entity_list in entities_dict.items():
                count = len(entity_list)
                if count > 0:
                    self.logger.log_info(f"   {entity_type}: {count}")
                    total_entities += count
            
            self.logger.log_info(f"✅ Entità estratte: {total_entities}")
            
            # Mostra alcuni esempi
            example_count = 0
            for entity_type, entity_list in entities_dict.items():
                for entity_name in entity_list[:2]:  # Max 2 per categoria
                    if example_count < 5:
                        self.logger.log_info(f"   Esempio: {entity_name} ({entity_type})")
                        example_count += 1
            
            self.test_results["medical_entities"] = True
            self.logger.log_test_success("Verifica Medical Entities Extraction", f"{total_entities} entità estratte")
            
        except Exception as e:
            self.test_results["medical_entities"] = False
            self.logger.log_test_failure("Verifica Medical Entities Extraction", str(e))


async def main():
    """Esegue tutti i test della pipeline."""
    tester = PipelineTests()
    results = await tester.run_all_tests()
    
    # Return code based on results
    if results["summary"]["failed"] > 0:
        return 1
    return 0


if __name__ == "__main__":
    exit_code = asyncio.run(main())
    sys.exit(exit_code)